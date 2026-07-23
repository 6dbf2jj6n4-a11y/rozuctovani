"""
Generovani dokumentu Karty najemce (.docx) - Priloha c. 1 ke Smlouve.

Na rozdil od Smlouvy (viz core/contract_generator.py) se dokument sklada
od nuly pomoci python-docx, nevyplnuje se existujici sablona - Karta
nema pevny puvodni vzor, jde o tabulkovy vypis dat Karty klienta
(plochy, klice rozuctovani sluzeb) plus podpisove pole.
"""
from decimal import Decimal

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.contract_generator import LANDLORD_REPRESENTATIVE, format_date_cz
from core.models import AllocationKey, ServicePoolItem

_CLASS_ORDER = [code for code, _ in ServicePoolItem.InvoiceClass.choices]


def _set_cell_text(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def _add_table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True)
    return table


def _fmt_num(value, decimals=2):
    if value is None:
        return "—"
    return f"{value:.{decimals}f}"


def generate_client_card_document(card, output_path):
    """Vygeneruje Kartu nájemce (Příloha č. 1) pro danou ClientCard a uloží
    do output_path (cesta nebo zapisovatelný stream, např. BytesIO)."""
    doc = docx.Document()

    header = doc.add_paragraph("Příloha č. 1")
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].bold = True

    doc.add_heading("KARTA NÁJEMCE", level=1)

    info = doc.add_paragraph()
    info.add_run("Klient: ").bold = True
    info.add_run(str(card.client))
    info.add_run("\n")
    info.add_run("Karta: ").bold = True
    info.add_run(card.description or f"Karta {card.client}")
    info.add_run("\n")
    info.add_run("Platnost od: ").bold = True
    info.add_run(format_date_cz(card.valid_from))
    if card.valid_to:
        info.add_run("\n")
        info.add_run("Platnost do: ").bold = True
        info.add_run(format_date_cz(card.valid_to))

    # --- Plochy ---
    doc.add_heading("Pronajaté plochy", level=2)
    units_table = _add_table(
        doc, ["Plocha", "Výměra (m²)", "Cena (Kč/m²/rok)", "Nájemné/rok (Kč)", "Nájemné/měsíc (Kč)"]
    )

    total_area = Decimal("0")
    total_year = Decimal("0")
    total_month = Decimal("0")
    for cu in card.card_units.select_related("unit__site"):
        area = cu.area_m2
        year_rent = (area * cu.rate_per_m2) if (area and cu.rate_per_m2) else None
        month_rent = cu.monthly_rent

        row = units_table.add_row().cells
        _set_cell_text(row[0], str(cu.unit) if cu.unit else "—")
        _set_cell_text(row[1], _fmt_num(area))
        _set_cell_text(row[2], _fmt_num(cu.rate_per_m2))
        _set_cell_text(row[3], _fmt_num(year_rent))
        _set_cell_text(row[4], _fmt_num(month_rent))

        total_area += area or Decimal("0")
        total_year += year_rent or Decimal("0")
        total_month += month_rent or Decimal("0")

    total_row = units_table.add_row().cells
    _set_cell_text(total_row[0], "Celkem", bold=True)
    _set_cell_text(total_row[1], _fmt_num(total_area), bold=True)
    _set_cell_text(total_row[2], "", bold=True)
    _set_cell_text(total_row[3], _fmt_num(total_year), bold=True)
    _set_cell_text(total_row[4], _fmt_num(total_month), bold=True)

    # --- Klíče, po sekcích dle třídy ---
    doc.add_heading("Klíče rozúčtování služeb", level=2)
    keys = list(
        card.allocation_keys
        .select_related("service_item", "meter")
        .order_by("service_item__invoice_class", "service_item__name")
    )
    class_labels = dict(ServicePoolItem.InvoiceClass.choices)
    type_labels = dict(AllocationKey.AllocationType.choices)

    for class_code in _CLASS_ORDER:
        class_keys = [k for k in keys if k.service_item.invoice_class == class_code]
        if not class_keys:
            continue
        doc.add_heading(class_labels[class_code], level=3)
        keys_table = _add_table(doc, ["Položka", "Typ výpočtu", "Hodnota", "Měřidlo"])
        for key in class_keys:
            row = keys_table.add_row().cells
            _set_cell_text(row[0], key.service_item.name)
            _set_cell_text(row[1], type_labels.get(key.allocation_type, key.allocation_type))
            _set_cell_text(row[2], _fmt_num(key.value, decimals=4) if key.value is not None else "—")
            _set_cell_text(row[3], str(key.meter) if key.meter else "—")

    # --- Podpisy ---
    doc.add_paragraph()
    doc.add_paragraph("_" * 30 + "\t\t\t" + "_" * 30)
    doc.add_paragraph(f"za Pronajímatele: {LANDLORD_REPRESENTATIVE}\t\tza Nájemce: {card.client}")

    doc.save(output_path if hasattr(output_path, "write") else str(output_path))
    return output_path
