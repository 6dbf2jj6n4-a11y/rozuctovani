"""
Generovani dokumentu Smlouvy (.docx) z sablony core/contract_templates/smlouva_template.docx.

Sablona pouziva viditelne placeholdery ve tvaru [NAZEV_PLACEHOLDERU] (napr.
[NÁZEV NÁJEMCE], [DATUM PODPISU]) na mistech, kam se dosazuji konkretni
udaje z DB - diky tomu jde sablonu otevrit primo ve Wordu a hned je jasne,
co je promenlive a co pevny text Smlouvy, aniz by to vypadalo jako uz
vyplnena konkretni smlouva. Generovani funguje tak, ze najde konkretni
run(y) obsahujici presne tento placeholder text a prepise jen jejich TEXT -
font/velikost/barva runu se nikde nevynucuje ani nemeni, zustava presne
takovy, jaky uz v sablone je. Kazda dosazena hodnota navic dostane zluty
zvyraznovac (viz _highlight), aby bylo v dokumentu na prvni pohled videt,
co se automaticky doplnilo a co je potreba pred odeslanim zkontrolovat.

Pri pridavani noveho placeholderu do sablony: musi byt v ramci celeho
dokumentu TEXTOVE UNIKATNI (marker/old se hleda jako prvni substring shoda
- viz _find_paragraph_index/_replace_run_text), jinak se dosadi na spatne
misto.

Blok Pronajimatele v zahlavi dokumentu (odstavce 6-13) je zamerne staticky
a NEzvyraznuje se - Pronajimatel (CALAMARI SE) je v teto aplikaci jeden
jediny a v sablone uz ma spravne udaje; dynamicky (a zvyrazneny) se
doplnuje jen nazev v zahlavi kazde stranky a jmeno zastupce v podpisovem
radku (viz nize) - tam, kde je uprava bezriziko (samostatny run).

Clanek 1 (Vymezeni predmetu a ucelu najmu) je pro kazdy areal jiny (vypis
jinych pozemku/budov) - jeho text se necte ze sablony, ale z
core.models.Site.lease_subject_text (viz _fill_article1). Pocet odstavcu
vycetu nemovitosti se tak muze mezi generovanimi lisit → cokoli za clankem 1
se MUSI dohledavat podle textu, ne podle pevneho indexu (na rozdil od bloku
najemce v zahlavi, ktery je pred clankem 1 a indexy tam zustavaji stabilni).
"""
import copy
from decimal import Decimal
from pathlib import Path

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as _RT
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

TEMPLATE_PATH = Path(__file__).resolve().parent / "contract_templates" / "smlouva_template.docx"

_MONTHS = [
    "ledna", "února", "března", "dubna", "května", "června",
    "července", "srpna", "září", "října", "listopadu", "prosince",
]

# Male, uzavrene mnozina soudu vedoucich obchodni rejstrik v CR - staci pro spravny
# gramaticky pad (7. pad/instrumental) v pripadech, ktere se v praxi objevi.
_COURT_INSTRUMENTAL_PREFIXES = {
    "krajský soud v": "krajským soudem v",
    "městský soud v praze": "městským soudem v Praze",
}


def _court_instrumental(court_name):
    """'Krajský soud v Ostravě' -> 'Krajským soudem v Ostravě'. Neznamy tvar se
    vrati beze zmeny - v generovanem dokumentu je pak potreba rucne zkontrolovat pad."""
    if not court_name:
        return ""
    lowered = court_name.strip().lower()
    for prefix, replacement in _COURT_INSTRUMENTAL_PREFIXES.items():
        if lowered.startswith(prefix):
            rest = court_name.strip()[len(prefix):]
            first_word, _, tail = replacement.partition(" ")
            return f"{first_word[0].upper()}{first_word[1:]} {tail}{rest}".rstrip()
    return court_name.strip()


def format_date_cz(d):
    if not d:
        return ""
    return f"{d.day}. {_MONTHS[d.month - 1]} {d.year}"


def format_czk(amount):
    if amount is None:
        return ""
    if isinstance(amount, Decimal):
        amount = int(amount)
    return f"{amount:,}".replace(",", " ") + " Kč"


def format_months(n):
    if n is None:
        return ""
    if n == 1:
        return f"{n} měsíc"
    if 2 <= n <= 4:
        return f"{n} měsíce"
    return f"{n} měsíců"


def get_landlord():
    """Vrati Klienta oznaceneho jako Pronajimatel (Client.is_landlord=True).

    Aplikace pocita s tim, ze takovy Klient je v databazi prave jeden -
    vynucuje se to na urovni admin formulare (Client.clean(), viz core/models.py),
    ne na urovni databaze, takze tato funkce prípadnou duplicitu neresí a
    jednoduse vezme prvniho nalezeneho."""
    from core.models import Client
    landlord = Client.objects.filter(is_landlord=True).first()
    if landlord is None:
        raise ValueError(
            "V databazi neni zadny Klient oznaceny jako Pronajimatel - "
            "oznac prislusneho Klienta prepinacem 'Pronajímatel' v adminu."
        )
    return landlord


def _format_address(obj):
    address = " ".join(p for p in (obj.street, obj.street_number) if p)
    if obj.zip_code or obj.city:
        address = f"{address}, {obj.zip_code} {obj.city}".strip(", ")
    return address


def contract_to_template_data(contract):
    """Sestavi `data` dict pro fill_contract_template() z instance Contract
    (a jejiho navazaneho Client) - sdileno mezi hromadnou akci v seznamu
    Smluv a tlacitkem generovani na detailu jedne Smlouvy."""
    client = contract.client
    landlord = get_landlord()

    return {
        "site_name": str(contract.site) if contract.site else "",
        "lease_subject_text": (contract.site.lease_subject_text if contract.site else "") or "",
        "landlord_name": landlord.name,
        "landlord_representative_name": landlord.representative_name,
        "client_name": client.name,
        "client_address": _format_address(client),
        "client_ico": client.ico,
        "client_dic": client.dic,
        "registry_court": client.registry_court,
        "registry_section": client.registry_section,
        "registry_insert": client.registry_insert,
        "representative_name": contract.representative_name,
        "representative_role": contract.representative_role,
        "invoicing_email": contract.invoicing_email,
        "signed_on": contract.signed_on,
        "valid_from": contract.valid_from,
        "valid_to": contract.valid_to,
        "notice_period_months": contract.notice_period_months,
        "insurance_amount_czk": contract.insurance_amount_czk,
        "deposit_czk": contract.deposit_czk,
        "inflation_increase_from": contract.inflation_increase_from,
    }


def _highlight(run):
    """Oznaci run zlutym zvyraznovacem - pouziva se na kazdou hodnotu
    dosazenou z databaze, aby bylo pred odeslanim smlouvy videt, co se
    doplnilo automaticky."""
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _set_paragraph_text_keep_format(paragraph, text, highlight=True):
    """Prepise cely text odstavce do jednoho runu - PONECHA font/velikost/barvu
    puvodniho prvniho runu (pokud existoval), jen font/velikost/barva se
    nevynucuje. Nehodi se pro odstavce s vice styly textu v jedne vete (tady
    zadny takovy mezi upravovanymi odstavci neni)."""
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
        target = runs[0]
    else:
        target = paragraph.add_run(text)
    if highlight:
        _highlight(target)


def _find_paragraph_index(paragraphs, marker):
    """Vrati index prvniho odstavce obsahujiciho `marker`, nebo vyhodi
    ValueError - pouziva se pro vse ZA clankem 1, protoze pocet jeho
    odstavcu je promenlivy (viz _fill_article1) a pevne indexy uz tam
    nejsou spolehlive."""
    for i, para in enumerate(paragraphs):
        if marker in para.text:
            return i
    raise ValueError(
        f"V šabloně smlouvy nenalezen odstavec obsahující {marker!r} - "
        f"zkontroluj core/contract_templates/smlouva_template.docx."
    )


def _split_and_highlight(run, old, new):
    """Nahradi PRVNI vyskyt `old` v textu runu za `new` a zvyrazni JEN tuto
    vlozenou cast - okolni puvodni text (pred/za `old` v ramci puvodniho
    runu) zustane beze zmeny stylu i bez zvyrazneni. Nutne proto, ze Word
    pri ukladani casto sloucí sousedni runy se stejnym formatovanim
    (napr. cela veta v jednom runu) - zvyraznit cely takovy run by obarvilo
    mnohem vic textu, nez jen dosazenou hodnotu. Puvodni run se rozdeli
    na az 3 runy se stejnym formatovanim (font/velikost/barva/atd.), jake
    mel puvodni cely run - zvyrazni se jen ten prostredni."""
    text = run.text
    pos = text.find(old)
    before, after = text[:pos], text[pos + len(old):]

    r_elem = run._element
    new_r = copy.deepcopy(r_elem)
    after_r = copy.deepcopy(r_elem)

    # Puvodni run mohl uz mit zvyrazneni (napr. kdyz nekdo v sablone rucne
    # oznacil celou vetu) - deepcopy by ho pak preneslo i do okolniho textu,
    # ktery zvyraznit nechceme, proto se explicitne vycisti na vsech 3 castech
    # pred tim, nez se zvyrazni jen ta prostredni (vlozena hodnota).
    run.text = before
    run.font.highlight_color = None

    new_run = Run(new_r, run._parent)
    new_run.text = new
    new_run.font.highlight_color = None

    after_run = Run(after_r, run._parent)
    after_run.text = after
    after_run.font.highlight_color = None

    r_elem.addnext(new_r)
    new_r.addnext(after_r)
    _highlight(new_run)

    if not before:
        r_elem.getparent().remove(r_elem)
    if not after:
        after_r.getparent().remove(after_r)


def _replace_run_text(paragraphs, marker, old, new, highlight=True):
    """Najde odstavec obsahujici `marker` (unikatni text pro dohledani spravneho
    mista - napr. cely puvodni datum, ne jen cislo, aby se nespletlo s jinym
    vyskytem stejne hodnoty jinde ve smlouve) a v nem run obsahujici `old`;
    v tomto runu nahradi `old` za `new`. Font/velikost/barva puvodniho textu
    se nemeni, jen se pripadne prida zvyraznovac na vlozenou hodnotu
    (viz _split_and_highlight - ne na cely run, ktery muze obsahovat
    i okolni nezmeneny text)."""
    idx = _find_paragraph_index(paragraphs, marker)
    para = paragraphs[idx]
    for run in para.runs:
        if old in run.text:
            if highlight:
                _split_and_highlight(run, old, new)
            else:
                run.text = run.text.replace(old, new)
            return
    raise ValueError(
        f"Marker {old!r} nalezen v odstavci {idx}, ale ne v samostatnem runu - "
        f"zkontroluj core/contract_templates/smlouva_template.docx."
    )


def _set_hyperlink(paragraph, new_url, new_text):
    """Prepise text a cil existujiciho hypertextoveho odkazu (w:hyperlink)
    v odstavci - pouziva se pro e-mailovou adresu pro elektronickou fakturaci.
    Bez `new_url` (prazdny e-mail) se odkaz cely odstrani - jinak by v dokumentu
    zustal "mrtvy" odkaz na puvodni (uz nespravnou) adresu ze sablony. Jinak
    puvodni vztah (relationship) na starou URL se ponecha nevyuzity (neskodi),
    misto prepisu se jen prida novy a hyperlink se prepoji na nej."""
    hyperlink_el = paragraph._p.find(qn("w:hyperlink"))
    if hyperlink_el is None:
        return
    if not new_url:
        hyperlink_el.getparent().remove(hyperlink_el)
        return
    t_el = hyperlink_el.find(".//" + qn("w:t"))
    if t_el is not None:
        t_el.text = new_text
    hl = hyperlink_el.find(".//" + qn("w:highlight"))
    if hl is None:
        rpr = hyperlink_el.find(".//" + qn("w:rPr"))
        if rpr is not None:
            hl = rpr.makeelement(qn("w:highlight"), {})
            rpr.append(hl)
    if hl is not None:
        hl.set(qn("w:val"), "yellow")
    new_r_id = paragraph.part.relate_to(new_url, _RT.HYPERLINK, is_external=True)
    hyperlink_el.set(qn("r:id"), new_r_id)


def _fill_article1(doc, lease_subject_text):
    """Dosadi text clanku 1 (Vymezeni predmetu a ucelu najmu) podle arealu.

    `lease_subject_text` je vicerádkový text (Site.lease_subject_text):
    prvni radek je uvodni veta (katastralni urad/LV/obec/k.u.), kazdy
    dalsi radek jeden pozemek/budova z vyctu. Uvodni odstavec sablony se
    prepise na miste; vycet pozemku/budov (puvodne pevny pocet odstavcu)
    se podle poctu radku bud zkrati (prebytecne odstavce smazou), nebo
    prodlouzi (posledni odrazkovy odstavec se naklonuje - sdili Word
    cislovani seznamu numId, takze klon zustane spravne formatovany
    i cislovany, vc. puvodniho fontu).

    Vraci pocet odstavcu, o ktery se cely dokument zmensil/zvetsil oproti
    puvodnimu stavu (kladne cislo = pribylo odstavcu) - volajici podle
    toho vi, ze indexy odstavcu ZA timto blokem uz nejsou spolehlive a
    musi je dohledavat podle textu.
    """
    paragraphs = doc.paragraphs
    # Presna shoda, ne substring - "Vymezení předmětu a účelu nájmu" se jako
    # PREFIX objevuje uz drive v Obsahu (napr. "Vymezení předmětu a účelu
    # nájmu\t3"), substring hledani (_find_paragraph_index) by tedy chybne
    # naslo radek Obsahu misto skutecneho nadpisu clanku.
    heading_idx = next(
        i for i, para in enumerate(paragraphs)
        if para.text.strip() == "Vymezení předmětu a účelu nájmu"
    )

    intro_idx = heading_idx + 1
    while not paragraphs[intro_idx].text.strip():
        intro_idx += 1

    stop_idx = intro_idx + 1
    while "Pronajímatel přenechává touto smlouvou" not in paragraphs[stop_idx].text:
        stop_idx += 1

    intro_para = paragraphs[intro_idx]
    bullet_paras = paragraphs[intro_idx + 1:stop_idx]

    lines = [line.strip() for line in (lease_subject_text or "").splitlines() if line.strip()]
    if not lines:
        _set_paragraph_text_keep_format(
            intro_para,
            "[DOPLNIT VYMEZENÍ PŘEDMĚTU NÁJMU PRO TENTO AREÁL - viz Areál v adminu]",
        )
        for bp in bullet_paras:
            bp._p.getparent().remove(bp._p)
        return -len(bullet_paras)

    intro_text, bullet_lines = lines[0], lines[1:]
    _set_paragraph_text_keep_format(intro_para, intro_text)

    n_existing = len(bullet_paras)
    n_needed = len(bullet_lines)

    for i in range(min(n_existing, n_needed)):
        _set_paragraph_text_keep_format(bullet_paras[i], bullet_lines[i])

    if n_needed > n_existing:
        template_xml = bullet_paras[-1]._p if bullet_paras else intro_para._p
        anchor = template_xml
        for i in range(n_existing, n_needed):
            new_xml = copy.deepcopy(template_xml)
            anchor.addnext(new_xml)
            anchor = new_xml
            new_para = Paragraph(new_xml, bullet_paras[-1]._parent if bullet_paras else intro_para._parent)
            _set_paragraph_text_keep_format(new_para, bullet_lines[i])
    elif n_needed < n_existing:
        for bp in bullet_paras[n_needed:]:
            bp._p.getparent().remove(bp._p)

    return n_needed - n_existing


def fill_contract_template(data, output_path, template_path=TEMPLATE_PATH):
    """
    `data` je dict s klici:
      site_name, lease_subject_text (viz core.models.Site.lease_subject_text),
      landlord_name, landlord_representative_name,
      client_name, client_address (jednoradkovy retezec),
      client_ico, client_dic (s nebo bez "CZ" prefixu), registry_court,
      registry_section, registry_insert, representative_name, representative_role,
      invoicing_email, signed_on (date), valid_from (date),
      valid_to (date/None - None = na dobu neurčitou, jinak na dobu určitou
      do tohoto data, viz cl. 7.1 a poznamka u is_fixed_term nize),
      notice_period_months (int), insurance_amount_czk (Decimal/int),
      deposit_czk (Decimal/int), inflation_increase_from (date).

    Chybejici/None hodnoty se dosadi jako prazdny retezec - vysledny dokument
    je pak potreba pred odeslanim zkontrolovat.
    """
    doc = docx.Document(str(template_path))

    client_name = data.get("client_name") or ""
    landlord_name = data.get("landlord_name") or ""

    # --- zahlavi dokumentu (nahoře na kazde strance): prvni run je vzdy kod/
    # nazev arealu; zbytek ("CALAMARI SE / nájemce") se dohledava podle textu,
    # ne podle pozice runu - pocet runu tam neni pevny (zavisi na tom, jak
    # sablonu naposledy nekdo v OK edit oval) a Pronajimatel se nezvyraznuje. ---
    header_para = doc.sections[0].header.paragraphs[0]
    header_runs = header_para.runs
    header_runs[0].text = data.get("site_name") or ""
    _highlight(header_runs[0])
    for run in header_runs:
        if "CALAMARI SE" in run.text:
            run.text = run.text.replace("CALAMARI SE", landlord_name)
        if "nájemce" in run.text:
            run.text = run.text.replace("nájemce", client_name)
            _highlight(run)

    # --- blok Najemce v zahlavi (odstavce 19-25, PRED clankem 1 - indexy
    # tedy zustavaji stabilni bez ohledu na delku vyctu nemovitosti v cl. 1).
    # Blok Pronajimatele (6-13) je zamerne staticky, viz modulovy docstring. ---
    paragraphs = doc.paragraphs
    paragraphs[19].runs[0].text = client_name
    _highlight(paragraphs[19].runs[0])
    paragraphs[20].runs[-1].text = data.get("client_address") or ""
    _highlight(paragraphs[20].runs[-1])
    paragraphs[21].runs[-1].text = data.get("client_ico") or ""
    _highlight(paragraphs[21].runs[-1])

    dic = (data.get("client_dic") or "").strip()
    dic_digits = dic[2:] if dic.upper().startswith("CZ") else dic
    paragraphs[22].runs[0].text = "DIČ: CZ" + dic_digits
    _highlight(paragraphs[22].runs[0])

    court = _court_instrumental(data.get("registry_court"))
    section = data.get("registry_section") or ""
    insert = data.get("registry_insert") or ""
    reg_runs = paragraphs[23].runs
    reg_runs[1].text = court
    _highlight(reg_runs[1])
    reg_runs[3].text = section
    _highlight(reg_runs[3])
    reg_runs[5].text = insert
    _highlight(reg_runs[5])

    rep_runs = paragraphs[24].runs
    rep_runs[1].text = data.get("representative_name") or ""
    _highlight(rep_runs[1])
    rep_runs[3].text = data.get("representative_role") or ""
    _highlight(rep_runs[3])

    email = data.get("invoicing_email") or ""
    _set_hyperlink(paragraphs[25], f"mailto:{email}" if email else None, email)

    # --- clanek 1: vlastni text podle arealu (meni pocet odstavcu dokumentu) ---
    _fill_article1(doc, data.get("lease_subject_text"))

    # --- cl. 7.1 (Trvani Smlouvy): sablona obsahuje DVE varianty tohoto
    # odstavce (na dobu neurcitou / na dobu urcitou), obe se stejnym
    # cislovanim (numId) - podle Contract.valid_to se ta nepouzita SMAZE
    # jeste PRED doplnovanim placeholderu nize, protoze obe varianty sdili
    # placeholder "[DATUM PLATNOSTI SMLOUVY]" a kdyby zbyly obe, doplnil by
    # se jen do prvni z nich (viz _find_paragraph_index). U smlouvy na dobu
    # urcitou navic mizi odst. 7.2 pism. b) "vypoved bez udani duvodu" (na
    # pevnem konci nedava smysl, viz zadani) a s nim i navazujici odkaz na
    # "pismene (B)" v puvodnim pism. c), ktere se tak stane novym pism. b). ---
    is_fixed_term = bool(data.get("valid_to"))
    paragraphs = doc.paragraphs
    if is_fixed_term:
        idx = _find_paragraph_index(paragraphs, "na dobu neurčitou")
    else:
        idx = _find_paragraph_index(paragraphs, "na dobu určitou")
    paragraphs[idx]._p.getparent().remove(paragraphs[idx]._p)

    if is_fixed_term:
        paragraphs = doc.paragraphs
        idx_b = _find_paragraph_index(paragraphs, "vypovědět i bez udání důvodu")
        paragraphs[idx_b]._p.getparent().remove(paragraphs[idx_b]._p)

        paragraphs = doc.paragraphs
        _replace_run_text(
            paragraphs, "písmene (B) nepoužije",
            " V\xa0případě výpovědi dle tohoto písmene se ustanovení písmene (B) nepoužije.", "",
            highlight=False,
        )

    # --- vse ZA clankem 1: dohledavat podle textu, ne podle indexu (viz docstring) ---
    paragraphs = doc.paragraphs
    _replace_run_text(
        paragraphs, "[DATUM ZVÝŠENÍ NÁJEMNÉHO]", "[DATUM ZVÝŠENÍ NÁJEMNÉHO]",
        format_date_cz(data.get("inflation_increase_from")),
    )
    paragraphs = doc.paragraphs
    _replace_run_text(
        paragraphs, "[POJISTNÁ ČÁSTKA]", "[POJISTNÁ ČÁSTKA]",
        format_czk(data.get("insurance_amount_czk")),
    )
    paragraphs = doc.paragraphs
    _replace_run_text(
        paragraphs, "[DATUM PLATNOSTI SMLOUVY]", "[DATUM PLATNOSTI SMLOUVY]",
        format_date_cz(data.get("valid_from")),
    )
    if is_fixed_term:
        paragraphs = doc.paragraphs
        _replace_run_text(
            paragraphs, "[DATUM UKONČENÍ SMLOUVY]", "[DATUM UKONČENÍ SMLOUVY]",
            format_date_cz(data.get("valid_to")),
        )
    if not is_fixed_term:
        # placeholder je jen v odst. 7.2 pism. b) "vypoved bez udani duvodu",
        # ktere u smlouvy na dobu urcitou vyse jiz zmizelo
        paragraphs = doc.paragraphs
        _replace_run_text(
            paragraphs, "[VÝPOVĚDNÍ LHŮTA]", "[VÝPOVĚDNÍ LHŮTA]",
            format_months(data.get("notice_period_months")),
        )
    paragraphs = doc.paragraphs
    _replace_run_text(
        paragraphs, "[VÝŠE KAUCE]", "[VÝŠE KAUCE]",
        format_czk(data.get("deposit_czk")),
    )
    paragraphs = doc.paragraphs
    _replace_run_text(
        paragraphs, "V\xa0Ostravě dne [DATUM PODPISU]", "[DATUM PODPISU]",
        format_date_cz(data.get("signed_on")),
    )
    paragraphs = doc.paragraphs
    _replace_run_text(
        paragraphs, "výpisu z obchodního rejstříku pořízeného ke dni", "[DATUM VÝPISU Z OR]",
        format_date_cz(data.get("signed_on")),
    )

    # --- podpisovy radek: jmena zastupcu pod carou a pod nimi nazvy
    # smluvnich stran (nadpis "Pronajímatel/Nájemce" a cara nad nimi jsou
    # staticky text sablony, viz modulovy docstring) ---
    paragraphs = doc.paragraphs
    names_idx = _find_paragraph_index(paragraphs, "\tIng. Daniel DAVID")
    names_runs = paragraphs[names_idx].runs
    names_runs[0].text = "\t" + (data.get("landlord_representative_name") or "")
    names_runs[-1].text = data.get("representative_name") or ""
    _highlight(names_runs[-1])

    company_runs = paragraphs[names_idx + 1].runs
    company_runs[0].text = company_runs[0].text.replace("CALAMARI SE", landlord_name)
    _split_and_highlight(company_runs[-1], "název nájemce", client_name)

    # python-docx prijima jak cestu (str/Path), tak zapisovatelny stream (napr. BytesIO)
    doc.save(output_path if hasattr(output_path, "write") else str(output_path))
    return output_path
